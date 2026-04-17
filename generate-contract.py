#!/usr/bin/env python3
import sys
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import os

def generate_contract(data):
    """Generiert einen Vertrag basierend auf Kundendaten"""
    
    # Lade den Template-Vertrag
    template_path = '/home/ubuntu/f26-calculator/client/public/template-vertrag.docx'
    doc = Document(template_path)
    
    # Ersetze Platzhalter in Absätzen
    for para in doc.paragraphs:
        if '[Name / Firma des Standortgebers]' in para.text:
            para.text = para.text.replace('[Name / Firma des Standortgebers]', data.get('gesetzlicherVertreter', ''))
        if '[Name / Firma des Aufstellers]' in para.text:
            para.text = para.text.replace('[Name / Firma des Aufstellers]', 'FitForFuture Energy Nord GmbH')
        if 'Anschrift: [●]' in para.text:
            if 'Standortgeber' in para.text or para.text.index('Anschrift:') < 50:
                para.text = para.text.replace('[●]', data.get('kundenAdresse', ''))
            else:
                para.text = para.text.replace('[●]', 'Melchiorstraße 26, 10179 Berlin')
    
    # Ersetze auch in Tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '[Name / Firma des Standortgebers]' in para.text:
                        para.text = para.text.replace('[Name / Firma des Standortgebers]', data.get('gesetzlicherVertreter', ''))
                    if '[●]' in para.text:
                        para.text = para.text.replace('[●]', data.get('kundenAdresse', ''))
    
    # Speichere als DOCX
    output_docx = f"/tmp/F26-Vertrag-{data.get('gesetzlicherVertreter', 'Kunde').replace(' ', '_')}.docx"
    doc.save(output_docx)
    
    # Konvertiere zu PDF
    output_pdf = output_docx.replace('.docx', '.pdf')
    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', '/tmp', output_docx], 
                      check=True, capture_output=True, timeout=30)
    except Exception as e:
        print(f"Fehler bei PDF-Konvertierung: {e}", file=sys.stderr)
        return None
    
    return output_pdf

if __name__ == '__main__':
    data = json.loads(sys.argv[1])
    pdf_path = generate_contract(data)
    if pdf_path:
        print(pdf_path)
    else:
        sys.exit(1)
